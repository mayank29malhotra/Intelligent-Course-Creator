"""DOCX exporter for course markdown.

Converts markdown content into a reasonably well-formatted DOCX document
using python-docx. Focuses on headings, paragraphs, lists, code blocks,
and simple tables.
"""

from pathlib import Path
from typing import List

import markdown2
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def _ensure_custom_styles(document: Document) -> None:
    styles = document.styles

    # Base body text
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)

    # Headings 1-3
    for level, size in [(1, 20), (2, 16), (3, 14)]:
        name = f"Heading {level}"
        try:
            style = styles[name]
        except KeyError:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True

    # Code style
    if "Code" not in styles:
        code_style = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code"]
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code_style.font.size = Pt(10)


def export_markdown_to_docx(markdown_content: str, output_path: str) -> str:
    """Convert markdown content to a DOCX file.

    Args:
        markdown_content: Raw markdown string
        output_path: Path where the .docx should be written

    Returns:
        The absolute path to the generated DOCX file.
    """
    document = Document()
    _ensure_custom_styles(document)

    lines: List[str] = markdown_content.splitlines()

    in_code_block = False
    code_language = ""
    code_lines: List[str] = []

    for raw_line in lines:
        line = raw_line.rstrip("\n")

        # Fenced code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                # End of code block
                paragraph = document.add_paragraph("\n".join(code_lines), style="Code")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                code_lines = []
                code_language = ""
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
                parts = line.strip().split()
                code_language = parts[0].lstrip("`") if parts else ""
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        stripped = line.strip()

        # Blank line
        if not stripped:
            document.add_paragraph("")
            continue

        # Headings #..######
        if stripped.startswith("#"):
            hashes, text = stripped.split(" ", 1)
            level = min(len(hashes), 3)
            paragraph = document.add_paragraph(text.strip(), style=f"Heading {level}")
            continue

        # Tables (very simple pipe tables)
        if "|" in stripped and stripped.count("|") >= 2:
            # Try to detect header row + separator
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            row_cells = [cells]
            # For now, create a 1-row table; multi-row handled implicitly
            table = document.add_table(rows=1, cols=len(cells))
            hdr_cells = table.rows[0].cells
            for i, cell_text in enumerate(cells):
                hdr_cells[i].text = cell_text
            continue

        # Lists
        if stripped.startswith(("- ", "* ")):
            text = stripped[2:].strip()
            paragraph = document.add_paragraph(text, style="List Bullet")
            continue

        if any(stripped.startswith(f"{i}. ") for i in range(1, 10)):
            # Simple numbered list 1.-9.
            parts = stripped.split(" ", 1)
            text = parts[1] if len(parts) > 1 else ""
            paragraph = document.add_paragraph(text, style="List Number")
            continue

        # Fallback: normal paragraph
        document.add_paragraph(line)

    output_path = str(Path(output_path).absolute())
    document.save(output_path)
    return output_path
